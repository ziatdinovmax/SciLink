"""DOCX → text, tables and figures, in document order.

``extract_text`` reads a .docx as ``"\\n".join(p.text for p in doc.paragraphs)``,
which silently discards two things a scientific document usually carries its
content in: every table, and every embedded figure. A white paper whose
methods sit in a table comes back with the methods missing, and nothing
reports the loss.

Ported from the peerpanel document loader (same author, MIT). Two decisions
from there are load-bearing and worth restating, because the obvious
implementations get both wrong:

**Document order, not `doc.paragraphs` then `doc.tables`.** python-docx
exposes paragraphs and tables as separate sequences, so the obvious join
appends every table at the end, detached from the heading it belongs under.
Walking ``document.element.body`` keeps a table where the reader sees it.

**Figures are marked in place AND returned.** Each embedded image emits a
``[Figure N]`` marker at its exact position in the text, and ``figures[N-1]``
is that image. DOCX has no page concept, so the marker is the only thing that
ties a figure to the caption or sentence it belongs to; passing the images
alone would leave the model unable to say which is which.
"""
from __future__ import annotations

import base64
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional


@dataclass
class Figure:
    """An embedded image, numbered to match its ``[Figure N]`` marker."""
    index: int          # 0-based; marker in the text is index + 1
    ext: str            # "png", "jpeg", ...
    bytes: bytes

    @property
    def mime(self) -> str:
        e = self.ext.lower()
        return "image/jpeg" if e in ("jpg", "jpeg") else f"image/{e}"

    def to_base64(self) -> str:
        return base64.b64encode(self.bytes).decode("ascii")


@dataclass
class Table:
    """A table rendered as GitHub-flavored markdown, in document order."""
    index: int
    markdown: str


@dataclass
class ParsedDocument:
    source: str
    text: str
    tables: List[Table] = field(default_factory=list)
    figures: List[Figure] = field(default_factory=list)

    def composed(self, max_chars: Optional[int] = None) -> str:
        """Text plus tables as one markdown string.

        Figures are deliberately NOT inlined here — they stay in ``figures``
        for a caller that can deliver images, and the ``[Figure N]`` markers
        already sitting in the text tell a text-only caller they exist.
        """
        parts = [self.text.strip()]
        for t in self.tables:
            parts.append(f"\n\n### Table {t.index + 1}\n\n{t.markdown}\n")
        out = "\n".join(parts)
        if max_chars is not None and len(out) > max_chars:
            out = out[:max_chars] + "\n\n[... truncated ...]"
        return out


def rows_to_markdown(rows: List[List[str]]) -> str:
    """Render a 2D list as a GitHub-flavored markdown table.

    Cells are flattened to one line and pipes escaped, so a multi-line or
    pipe-bearing cell cannot break the table's shape. Short rows are padded
    to the widest row rather than dropped.
    """
    if not rows:
        return ""
    cleaned = [[(c or "").strip().replace("\n", " ").replace("|", "\\|")
                for c in row] for row in rows]
    width = max(len(r) for r in cleaned)
    cleaned = [r + [""] * (width - len(r)) for r in cleaned]
    header, body = cleaned[0], cleaned[1:]
    lines = ["| " + " | ".join(header) + " |",
             "| " + " | ".join(["---"] * width) + " |"]
    lines += ["| " + " | ".join(r) + " |" for r in body]
    return "\n".join(lines)


def _paragraph_to_text(p, *, fig_counter_start: int):
    """Walk a paragraph's XML in order -> (text_with_markers, image_rIds).

    Reading ``p.text`` would be shorter and would lose three things: tabs and
    line breaks (they are elements, not characters), and the POSITION of any
    embedded image, which is exactly what makes a figure attributable to its
    caption.
    """
    from docx.oxml.ns import qn

    parts: List[str] = []
    rids: List[str] = []
    counter = fig_counter_start
    for el in p._element.iter():
        tag = el.tag
        if tag == qn("w:t"):
            if el.text:
                parts.append(el.text)
        elif tag == qn("w:tab"):
            parts.append("\t")
        elif tag == qn("w:br"):
            parts.append("\n")
        elif tag == qn("a:blip"):
            rid = el.get(qn("r:embed"))
            if rid:
                parts.append(f" [Figure {counter}] ")
                rids.append(rid)
                counter += 1
    return "".join(parts), rids


def load_docx(path, *, include_figures: bool = True) -> ParsedDocument:
    """Read a .docx into text, tables and figures, in document order."""
    import docx  # python-docx

    p = Path(path)
    doc = docx.Document(str(p))

    text_parts: List[str] = []
    tables: List[Table] = []
    rid_order: List[str] = []       # image relationship ids, document order

    body = doc.element.body
    para_map = {id(x._element): x for x in doc.paragraphs}
    tbl_map = {id(x._element): x for x in doc.tables}

    for child in body.iterchildren():
        para = para_map.get(id(child))
        if para is not None:
            text, rids = _paragraph_to_text(
                para, fig_counter_start=len(rid_order) + 1)
            if text.strip() or rids:
                text_parts.append(text)
            rid_order.extend(rids)
            continue

        tbl = tbl_map.get(id(child))
        if tbl is not None:
            # Cells go through the same walker, so an image inside a table
            # keeps its place in the global figure numbering.
            rows: List[List[str]] = []
            for row in tbl.rows:
                cells: List[str] = []
                for cell in row.cells:
                    chunks: List[str] = []
                    for cp in cell.paragraphs:
                        text, rids = _paragraph_to_text(
                            cp, fig_counter_start=len(rid_order) + 1)
                        if text:
                            chunks.append(text)
                        rid_order.extend(rids)
                    cells.append(" ".join(chunks))
                rows.append(cells)
            md = rows_to_markdown(rows)
            if md:
                idx = len(tables)
                tables.append(Table(index=idx, markdown=md))
                # Leave a marker where the table stood, so its position in the
                # narrative survives even though the table renders below.
                text_parts.append(f"[Table {idx + 1}]")

    figures: List[Figure] = []
    if include_figures:
        rels = doc.part.rels
        for rid in rid_order:
            rel = rels.get(rid)
            if rel is None or "image" not in getattr(rel, "reltype", ""):
                continue
            try:
                blob = rel.target_part.blob
                ctype = getattr(rel.target_part, "content_type", "image/png")
                figures.append(Figure(index=len(figures),
                                      ext=ctype.split("/")[-1], bytes=blob))
            except Exception as e:  # noqa: BLE001 - one bad image, not a failure
                logging.debug(f"Skipping unreadable image {rid}: {e}")

    return ParsedDocument(source=str(p), text="\n\n".join(text_parts),
                          tables=tables, figures=figures)


def count_docx_figures(path) -> int:
    """How many embedded images a .docx holds, without decoding them.

    Lets a caller decide whether attaching figures is affordable BEFORE
    paying for them — for DOCX the count is exact, so there is no need to
    estimate it from document length.
    """
    import zipfile
    try:
        with zipfile.ZipFile(str(path)) as z:
            xml = z.read("word/document.xml").decode("utf-8", "ignore")
        return xml.count("<a:blip")
    except Exception as e:  # noqa: BLE001 - a count is advisory
        logging.debug(f"Figure count failed for {path}: {e}")
        return 0
