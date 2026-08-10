"""Unified lightweight document text extraction (Tier 1).

``extract_text`` turns a single document into flat text — for callers that
want a handful of documents read straight into context, with no chunking,
embeddings, or retrieval. The heavyweight ingestion/retrieval path lives in
``scilink.knowledge``.
"""

import os
from pathlib import Path
from typing import Any, Dict, Union

from .pdf_parser import _extract_pdf_blocks, _assemble_flat_text
from .ocr import DEFAULT_OCR_DPI, MAX_OCR_PAGES

# Small LRU over full extractions of the expensive formats —
# see the cache comment in extract_text. 8 documents is ample
# for any session's working set; texts are ~100 KB each.
from collections import OrderedDict
_EXTRACT_CACHE: "OrderedDict[tuple, dict]" = OrderedDict()
_EXTRACT_CACHE_MAX = 8



def extract_text(path: Union[str, Path], max_pages: int = None,
                 ocr_model: Any = None, ocr_dpi: int = DEFAULT_OCR_DPI,
                 max_ocr_pages: int = MAX_OCR_PAGES) -> Dict[str, Any]:
    """Extract plain text (and markdown tables, for PDFs) from a document.

    Supports text-like (``.pdf``, ``.docx``, ``.md``, ``.txt``, ``.json``,
    ``.yaml``/``.yml``) and tabular (``.csv``, ``.xlsx``/``.xls``) documents.
    Returns a dict with ``text``, ``n_chars`` and a format-specific count
    (``n_pages`` for PDFs, ``n_paragraphs`` for DOCX). PDFs also report
    ``n_ocr_pages`` — pages transcribed via the vision-OCR fallback.
    Tabular files are previewed via the adaptive Excel parser — small files
    yield the full table as Markdown, large files a statistical summary; an
    auto-detected sibling JSON metadata file (e.g. ``data.json`` next to
    ``data.xlsx``) enriches the preview with title / objective / column
    definitions.

    No truncation is applied — any length cap is the caller's policy.

    Args:
        path: Path to the document.
        max_pages: For PDFs only — stop after this many pages and skip table
            extraction. A lightweight mode for previews/probes; ``n_pages``
            still reports the true total.
        ocr_model: Optional vision LLM (any object with ``generate_content``).
            When given, scanned/sparse PDF pages are transcribed via OCR;
            without it, such pages simply yield their (sparse) text.
        ocr_dpi: Render resolution for OCR'd pages.
        max_ocr_pages: Cap on the number of pages sent to vision-OCR.

    Raises:
        ValueError: For an unsupported extension.
        ImportError: For a ``.docx`` when ``python-docx`` is not installed.
    """
    path = Path(path)
    ext = path.suffix.lower()
    info: Dict[str, Any] = {}

    # Extraction cache for the expensive formats. Windowed readers
    # (planning read_file, meta view_document, analysis read_document)
    # legitimately call several times per document — without this, every
    # window re-parsed the whole PDF and re-ran vision-OCR on the same
    # scanned pages (a 200-DPI vision-LLM call per window, observed
    # live). Keyed on (path, mtime, size) plus every parameter that
    # changes the output; a rewritten file misses and re-extracts.
    # $SCILINK_EXTRACT_CACHE=0 disables.
    cache_key = None
    if ext in (".pdf", ".docx") and os.environ.get(
            "SCILINK_EXTRACT_CACHE", "").strip().lower() not in (
            "0", "false", "off", "no"):
        try:
            st = path.stat()
            cache_key = (str(path.resolve()), st.st_mtime_ns, st.st_size,
                         max_pages, bool(ocr_model), ocr_dpi, max_ocr_pages)
            hit = _EXTRACT_CACHE.get(cache_key)
            if hit is not None:
                _EXTRACT_CACHE.move_to_end(cache_key)
                return dict(hit)
        except OSError:
            cache_key = None

    if ext == ".pdf":
        page_texts, table_chunks, n_pages, ocr_pages = _extract_pdf_blocks(
            str(path), max_pages=max_pages, ocr_model=ocr_model,
            ocr_dpi=ocr_dpi, max_ocr_pages=max_ocr_pages,
        )
        text = _assemble_flat_text(page_texts, table_chunks, ocr_pages)
        info["n_pages"] = n_pages
        info["n_ocr_pages"] = len(ocr_pages)
    elif ext == ".docx":
        try:
            import docx  # noqa: F401 - checked here for a clear error
        except ImportError as e:
            raise ImportError(
                "Reading .docx documents requires python-docx "
                "(`pip install python-docx`)."
            ) from e
        # Tables were silently dropped here: paragraphs and tables are
        # separate sequences in python-docx, so joining p.text lost every
        # table in the document with nothing reporting the loss. The loader
        # walks the body in order instead, so a table lands where the reader
        # sees it. Figures are located and marked ([Figure N]) but not
        # returned on this path — extract_document() carries the bytes for
        # callers that can deliver images.
        from .docx_document import count_docx_figures, load_docx
        d = load_docx(path, include_figures=False)
        import docx as _docx
        info["n_paragraphs"] = len(_docx.Document(str(path)).paragraphs)
        info["n_tables"] = len(d.tables)
        info["n_figures"] = count_docx_figures(path)
        text = d.composed()
    elif ext in (".md", ".txt", ".json", ".yaml", ".yml"):
        text = path.read_text(errors="replace")
    elif ext in (".csv", ".xlsx", ".xls"):
        from .excel_parser import parse_adaptive_excel
        # Auto-discover a sibling JSON metadata file (e.g. data.xlsx +
        # data.json with title / objective / column_definitions) — the same
        # convention parse_adaptive_excel uses elsewhere in the codebase.
        sidecar = path.with_suffix(".json")
        chunks = parse_adaptive_excel(
            str(path),
            context_path=str(sidecar) if sidecar.exists() else None,
        )
        if chunks:
            summary = next(
                (c for c in chunks if c["metadata"].get("content_type")
                 in ("dataset_summary", "dataset_package")),
                chunks[0],
            )
            text = summary["text"]
        else:
            text = ""
    else:
        raise ValueError(
            f"Unsupported document type '{ext}' — extract_text handles "
            f".pdf, .docx, .md, .txt, .json, .yaml/.yml, "
            f".csv and .xlsx/.xls."
        )

    text = text.strip()
    info["text"] = text
    info["n_chars"] = len(text)
    if cache_key is not None:
        _EXTRACT_CACHE[cache_key] = dict(info)
        _EXTRACT_CACHE.move_to_end(cache_key)
        while len(_EXTRACT_CACHE) > _EXTRACT_CACHE_MAX:
            _EXTRACT_CACHE.popitem(last=False)
    return info


def extract_document(path: Union[str, Path], **kwargs):
    """Like :func:`extract_text`, but keeps figures as image bytes.

    Returns a ``ParsedDocument`` (text with ``[Figure N]`` / ``[Table N]``
    markers, tables as markdown, figures as bytes) for ``.docx``. Every other
    format returns a ``ParsedDocument`` carrying only the text
    :func:`extract_text` produces, so a caller can treat all formats
    uniformly and simply find ``figures`` empty.

    Separate from ``extract_text`` on purpose: images are worth real tokens,
    so a caller opts into carrying them rather than having every existing
    text consumer start paying for them.
    """
    from .docx_document import ParsedDocument, load_docx

    path = Path(path)
    if path.suffix.lower() == ".docx":
        return load_docx(path, include_figures=True)
    info = extract_text(path, **kwargs)
    return ParsedDocument(source=str(path), text=info.get("text", ""))
