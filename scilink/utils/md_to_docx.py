"""Markdown → .docx for one-page technical memos.

A deliberately small converter for the memo shape produced by
``write_technical_document(style="memo")``: a title, a header block of
bold-labelled lines, ``## `` sections, paragraphs with ``**bold**`` runs,
``- `` bullets and simple pipe tables. It is not a general markdown
renderer — a white paper keeps its PDF twin; the memo additionally gets
the .docx that is what actually circulates.

Deterministic and LLM-free, like ``md_to_pdf``.
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import List, Optional, Tuple

_INLINE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def _add_runs(par, text: str) -> None:
    """Emit ``text`` into ``par`` honouring **bold**, *italic*, `code`."""
    for tok in _INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = par.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            r = par.add_run(tok[1:-1]); r.italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = par.add_run(tok[1:-1]); r.font.name = "Consolas"
        else:
            par.add_run(tok)


def _split_table(lines: List[str]) -> Optional[List[List[str]]]:
    rows = []
    for ln in lines:
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{2,}:?", c) for c in cells if c):
            continue  # separator row
        rows.append(cells)
    return rows or None


def markdown_to_docx(md_path: Path, label: str = "TECHNICAL MEMO",
                     docx_path: Optional[Path] = None) -> Path:
    """Render ``md_path`` to a .docx twin (same stem) and return its path.

    Layout: US Letter, 1-inch margins, 11-pt body; a bold small-caps
    ``label`` line above the title; ``## `` headings as Heading 1;
    ``**Bold.**`` lead-ins preserved as bold runs.
    """
    from docx import Document
    from docx.shared import Pt, Inches

    md_path = Path(md_path)
    out = Path(docx_path) if docx_path else md_path.with_suffix(".docx")
    text = md_path.read_text(encoding="utf-8", errors="replace")

    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    for side in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, side, Inches(1))
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)

    lines = text.splitlines()
    i, first_h1_done = 0, False
    para_buf: List[str] = []

    def flush_para():
        nonlocal para_buf
        if para_buf:
            p = doc.add_paragraph()
            _add_runs(p, " ".join(s.strip() for s in para_buf))
            para_buf = []

    while i < len(lines):
        ln = lines[i]
        s = ln.strip()
        if not s:
            flush_para(); i += 1; continue
        if s.startswith("# ") and not first_h1_done:
            flush_para()
            lab = doc.add_paragraph(); r = lab.add_run(label); r.bold = True
            lab.paragraph_format.space_after = Pt(0)
            t = doc.add_paragraph(); r = t.add_run(s[2:].strip())
            r.bold = True; r.font.size = Pt(15)
            first_h1_done = True
            i += 1; continue
        m = re.match(r"^(#{1,6})\s+(.*)$", s)
        if m:
            flush_para()
            level = 1 if len(m.group(1)) <= 2 else min(len(m.group(1)) - 1, 3)
            doc.add_heading(m.group(2).strip(), level=level)
            i += 1; continue
        if s.startswith("|"):
            flush_para()
            j = i
            while j < len(lines) and lines[j].strip().startswith("|"):
                j += 1
            rows = _split_table(lines[i:j])
            if rows:
                ncol = max(len(r) for r in rows)
                tbl = doc.add_table(rows=len(rows), cols=ncol)
                tbl.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci in range(ncol):
                        cell = tbl.cell(ri, ci)
                        cell.text = ""
                        _add_runs(cell.paragraphs[0],
                                  row[ci] if ci < len(row) else "")
                        if ri == 0:
                            for r in cell.paragraphs[0].runs:
                                r.bold = True
                doc.add_paragraph()
            i = j; continue
        if re.match(r"^[-*+]\s+", s):
            flush_para()
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, re.sub(r"^[-*+]\s+", "", s))
            i += 1; continue
        if re.match(r"^\d+[.)]\s+", s):
            flush_para()
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, re.sub(r"^\d+[.)]\s+", "", s))
            i += 1; continue
        if s.startswith("!["):  # images: skip in a memo docx
            flush_para(); i += 1; continue
        # A header-block line ("**Purpose:** ...", "*subtitle*", "**CORE RULE** ...")
        # stands alone; ordinary prose lines are joined into one paragraph.
        if re.match(r"^(\*\*[^*]+\*\*|\*[^*]+\*)", s) and (
                s.startswith("**CORE") or re.match(r"^\*\*[A-Za-z ]+:\*\*", s)
                or (s.startswith("*") and not s.startswith("**") and s.endswith("*"))):
            flush_para()
            p = doc.add_paragraph(); _add_runs(p, s)
            p.paragraph_format.space_after = Pt(2)
            i += 1; continue
        para_buf.append(s)
        i += 1
    flush_para()
    doc.save(out)
    return out
