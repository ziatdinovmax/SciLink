"""view_document decides whether figures are worth the context.

Figures are worth real tokens, so attaching every one unconditionally would
tax every document read. For DOCX the embedded images can be counted exactly
without decoding them, so the decision is made on the true cost before paying
it: a short note carrying a few diagrams gets them, an atlas of forty reports
its count instead — and says so, rather than leaving the agent to wonder why
the [Figure N] markers point at nothing.
"""

import json
import tempfile
from pathlib import Path
from types import SimpleNamespace

import pytest

docx = pytest.importorskip("docx")
PIL = pytest.importorskip("PIL")

from scilink.agents.meta_agent.meta_orchestrator_tools import (  # noqa: E402
    MetaOrchestratorTools)


def _doc(path, n_figures, with_table=False):
    from docx.shared import Inches
    from PIL import Image
    img = path.parent / "i.png"
    Image.new("RGB", (40, 30), "navy").save(img)

    d = docx.Document()
    d.add_paragraph("Body text.")
    if with_table:
        t = d.add_table(rows=2, cols=2)
        t.rows[0].cells[0].text = "k"
        t.rows[0].cells[1].text = "v"
    for i in range(n_figures):
        d.add_paragraph(f"Caption {i + 1}:")
        d.add_picture(str(img), width=Inches(0.5))
    d.save(path)
    return path


def _tools():
    t = MetaOrchestratorTools.__new__(MetaOrchestratorTools)
    t.orch = SimpleNamespace(model=None)
    cap = {}
    t._register_tool = (
        lambda func, name, description, parameters, required=None:
        cap.update({name: func}))
    MetaOrchestratorTools._register_all_tools(t)
    return cap


@pytest.fixture(scope="module")
def tmp():
    return Path(tempfile.mkdtemp(prefix="viewdoc_"))


def _call(cap, path, **kw):
    return json.loads(cap["view_document"]([str(path)], **kw))


def test_a_few_figures_ride_along_automatically(tmp):
    cap = _tools()
    out = _call(cap, _doc(tmp / "few.docx", 3))
    assert out["status"] == "success"
    assert len(out.get("images_base64", [])) == 3
    assert "figures_not_attached" not in out


def test_a_figure_heavy_document_reports_instead_of_attaching(tmp):
    cap = _tools()
    out = _call(cap, _doc(tmp / "many.docx", 20))
    assert "images_base64" not in out
    note = out.get("figures_not_attached", "")
    assert "20 embedded figure" in note, note
    assert "figures='on'" in note, "the agent must be told how to get them"


def test_forcing_on_raises_the_limit_rather_than_nudging_it(tmp):
    """'on' used to stop at the auto threshold, so it quietly meant
    'on, up to 8'. Caught live, where the agent reported getting 8 of 20."""
    out = _call(_tools(), tmp / "many.docx", figures="on")
    n = len(out.get("images_base64", []))
    assert n > 8, f"'on' still capped at the auto threshold ({n})"
    assert n == 20, "all 20 fit under the hard ceiling and should arrive"


def test_partial_delivery_is_stated_not_left_to_be_inferred(tmp):
    """Past even the forced ceiling, some figures cannot be delivered.
    Unsaid, the reply looks complete while the remaining [Figure N] markers
    point at nothing."""
    out = _call(_tools(), _doc(tmp / "atlas.docx", 30), figures="on")
    n = len(out["images_base64"])
    assert n < 30, "the hard ceiling did not bound the payload"
    note = out.get("figures_attached", "")
    assert f"{n} of 30" in note, note
    assert "not delivered" in note


def test_complete_delivery_says_nothing_extra(tmp):
    """No note when everything arrived — silence means complete."""
    out = _call(_tools(), _doc(tmp / "three.docx", 3), figures="on")
    assert len(out["images_base64"]) == 3
    assert "figures_attached" not in out


def test_off_attaches_nothing_and_says_why(tmp):
    cap = _tools()
    out = _call(cap, tmp / "few.docx", figures="off")
    assert "images_base64" not in out
    assert "figures='off'" in out.get("figures_not_attached", "")


def test_a_document_without_figures_says_nothing_about_them(tmp):
    cap = _tools()
    out = _call(cap, _doc(tmp / "none.docx", 0))
    assert "images_base64" not in out
    assert "figures_not_attached" not in out, "no figures is not a warning"


def test_tables_are_read_regardless_of_the_figure_policy(tmp):
    cap = _tools()
    p = _doc(tmp / "tbl.docx", 0, with_table=True)
    out = _call(cap, p, figures="off")
    assert out["documents"][0]["n_tables"] == 1
    assert "| k | v |" in out["documents"][0]["text"]


def test_the_cap_counts_across_all_requested_documents(tmp):
    """Three documents of four figures each is twelve — over the limit."""
    cap = _tools()
    paths = [str(_doc(tmp / f"multi{i}.docx", 4)) for i in range(3)]
    out = json.loads(cap["view_document"](paths))
    assert "images_base64" not in out
    assert "12 embedded figure" in out.get("figures_not_attached", "")


def test_attached_images_are_valid_base64_png(tmp):
    import base64
    cap = _tools()
    out = _call(cap, tmp / "few.docx", figures="on")
    raw = base64.b64decode(out["images_base64"][0])
    assert raw.startswith(b"\x89PNG"), "not a decodable image"


if __name__ == "__main__":
    raise SystemExit(pytest.main([__file__, "-q"]))


def _corrupt_one_figure(tmp, name="broken.docx"):
    """A .docx whose SECOND image is corrupt but still declares image/png —
    the case a mime-only check cannot catch, since the type is a claim the
    bytes do not honour."""
    import zipfile
    from docx.shared import Inches
    from PIL import Image
    src = tmp / "_src.docx"
    d = docx.Document()
    for colour in ("seagreen", "crimson", "royalblue"):     # distinct, or
        img = tmp / f"{colour}.png"                          # python-docx
        Image.new("RGB", (40, 30), colour).save(img)         # dedupes them
        d.add_paragraph(f"{colour}:")
        d.add_picture(str(img), width=Inches(0.4))
    d.save(src)

    out = tmp / name
    zin, zout = zipfile.ZipFile(src), zipfile.ZipFile(out, "w")
    media = sorted(n for n in zin.namelist() if n.startswith("word/media/"))
    for n in zin.namelist():
        data = zin.read(n)
        if n == media[1]:
            data = b"\x89PNG\r\n\x1a\n" + b"\x00" * 400
        zout.writestr(n, data)
    zout.close(); zin.close()
    return out


def test_an_undeliverable_figure_is_named_not_silently_missing(tmp):
    out = _call(_tools(), _corrupt_one_figure(tmp), figures="on")
    notes = out.get("figures_undisplayable", [])
    assert notes, "a broken figure vanished without explanation"
    assert any("[Figure 2]" in n for n in notes), notes
    assert any("could not be decoded" in n for n in notes), notes


def test_the_good_figures_still_arrive_around_a_broken_one(tmp):
    out = _call(_tools(), _corrupt_one_figure(tmp, "broken2.docx"), figures="on")
    assert len(out["images_base64"]) == 2, "one bad image took the others down"


def test_a_broken_figure_does_not_renumber_the_rest(tmp):
    """If figure 2 were dropped, [Figure 3] in the text would point at the
    image delivered second — the mis-attribution the markers prevent."""
    from scilink.parsers.extract import extract_document
    figs = extract_document(_corrupt_one_figure(tmp, "broken3.docx")).figures
    assert [f.index for f in figs] == [0, 1, 2]
    assert [f.deliverable for f in figs] == [True, False, True]
