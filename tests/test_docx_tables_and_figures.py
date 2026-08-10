"""A Word document's tables and figures must survive being read.

`extract_text` read a .docx as `"\\n".join(p.text for p in doc.paragraphs)`.
python-docx exposes paragraphs and tables as separate sequences, so that join
dropped every table in the document — silently, with nothing reporting the
loss — and embedded figures were invisible entirely. A white paper whose
methods sit in a table came back with the methods missing.

The loader walks the document body in order, so a table lands where the reader
sees it, and marks each embedded image `[Figure N]` at its exact position so a
figure can be tied to the caption it belongs to. DOCX has no page concept;
the marker is the only thing that carries that association.
"""

import re
import tempfile
from pathlib import Path

import pytest

from scilink.parsers.docx_document import (
    count_docx_figures, load_docx, rows_to_markdown)
from scilink.parsers.extract import extract_document, extract_text

docx = pytest.importorskip("docx")
PIL = pytest.importorskip("PIL")


@pytest.fixture(scope="module")
def doc_with_everything():
    from docx.shared import Inches
    from PIL import Image

    tmp = Path(tempfile.mkdtemp(prefix="docxfix_"))
    img = tmp / "i.png"
    Image.new("RGB", (60, 40), "navy").save(img)

    d = docx.Document()
    d.add_heading("Design Note", 0)
    d.add_paragraph("Intro before any structure.")
    t = d.add_table(rows=3, cols=3)
    for r, row in enumerate([["Channel", "Observable", "Rate"],
                             ["MZI", "phase | mass", "1 fps"],
                             ["Raman", "speciation", "0.2 fps"]]):
        for c, v in enumerate(row):
            t.rows[r].cells[c].text = v
    d.add_paragraph("Text AFTER the table.")
    d.add_paragraph("The layout is shown here:")
    d.add_picture(str(img), width=Inches(1))
    d.add_paragraph("and the flow path here:")
    d.add_picture(str(img), width=Inches(1))
    p = tmp / "everything.docx"
    d.save(p)
    return p


# ------------------------------------------------------------------ tables


def test_table_content_is_no_longer_dropped(doc_with_everything):
    """The regression: every cell used to vanish."""
    text = extract_text(doc_with_everything)["text"]
    for cell in ("Channel", "Observable", "Rate", "MZI", "Raman",
                 "speciation", "0.2 fps"):
        assert cell in text, f"{cell!r} was dropped"


def test_a_table_lands_where_the_reader_sees_it(doc_with_everything):
    """Joining paragraphs then tables would append it after the prose."""
    text = extract_text(doc_with_everything)["text"]
    assert text.index("[Table 1]") < text.index("Text AFTER the table.")


def test_prose_around_the_table_survives(doc_with_everything):
    text = extract_text(doc_with_everything)["text"]
    assert "Intro before any structure." in text
    assert "Text AFTER the table." in text


def test_a_pipe_in_a_cell_cannot_break_the_table():
    """An unescaped pipe would add a phantom column on that row only."""
    md = rows_to_markdown([["a", "b"], ["x | y", "z"]])
    assert r"x \| y" in md
    # Count DELIMITERS — an escaped pipe is content, not a cell boundary.
    delims = [len(re.findall(r"(?<!\\)\|", line)) for line in md.splitlines()]
    assert delims == [3, 3, 3], delims


def test_ragged_rows_are_padded_not_dropped():
    md = rows_to_markdown([["a", "b", "c"], ["only one"]])
    assert len(md.splitlines()) == 3          # header, rule, one body row
    assert md.splitlines()[-1].count("|") == 4


def test_counts_are_reported(doc_with_everything):
    info = extract_text(doc_with_everything)
    assert info["n_tables"] == 1
    assert info["n_figures"] == 2


# ----------------------------------------------------------------- figures


def test_figures_are_marked_in_place_and_returned(doc_with_everything):
    d = extract_document(doc_with_everything)
    assert len(d.figures) == 2
    # The marker must sit with the sentence that introduces it.
    assert d.text.index("layout is shown here") < d.text.index("[Figure 1]")
    assert d.text.index("[Figure 1]") < d.text.index("flow path here")
    assert d.text.index("flow path here") < d.text.index("[Figure 2]")


def test_marker_numbering_indexes_the_figure_list(doc_with_everything):
    """`figures[N-1]` must BE the image `[Figure N]` refers to."""
    d = extract_document(doc_with_everything)
    for n, fig in enumerate(d.figures, start=1):
        assert f"[Figure {n}]" in d.text
        assert fig.index == n - 1
        assert fig.bytes and fig.mime.startswith("image/")


def test_figures_can_be_counted_without_decoding(doc_with_everything):
    assert count_docx_figures(doc_with_everything) == 2


def test_the_text_path_does_not_carry_image_bytes(doc_with_everything):
    """extract_text stays cheap: markers yes, pixels no."""
    d = load_docx(doc_with_everything, include_figures=False)
    assert d.figures == []
    assert "[Figure 1]" in d.text


def test_base64_round_trips(doc_with_everything):
    import base64
    d = extract_document(doc_with_everything)
    assert base64.b64decode(d.figures[0].to_base64()) == d.figures[0].bytes


# ------------------------------------------------------------- other formats


def test_non_docx_formats_return_an_empty_figure_list(tmp_path):
    """A uniform shape so a caller need not branch on extension."""
    md = tmp_path / "n.md"
    md.write_text("# Note\n\nbody\n")
    d = extract_document(md)
    assert d.figures == [] and d.tables == []
    assert "body" in d.text


def test_a_document_with_neither_is_unchanged(tmp_path):
    d = docx.Document()
    d.add_heading("Plain", 0)
    d.add_paragraph("No tables, no figures.")
    p = tmp_path / "plain.docx"
    d.save(p)

    info = extract_text(p)
    assert info["n_tables"] == 0 and info["n_figures"] == 0
    assert info["text"].strip() == "Plain\n\nNo tables, no figures."


if __name__ == "__main__":
    raise SystemExit(pytest.main([__file__, "-q"]))


# --------------------------------------------------- image format handling


def test_bmp_and_tiff_are_converted_rather_than_shipped(tmp_path):
    """A Word document carries scanned figures as BMP/TIFF routinely. The
    delivery layer guesses mime from magic bytes and falls back to PNG, so
    shipping one means a BMP arrives labelled PNG and breaks the request."""
    from docx.shared import Inches
    from PIL import Image
    from scilink.parsers.docx_document import SUPPORTED_IMAGE_MIMES

    d = docx.Document()
    for fmt, ext in [("PNG", "png"), ("BMP", "bmp"), ("TIFF", "tif")]:
        img = tmp_path / f"i.{ext}"
        Image.new("RGB", (40, 30), "orange").save(img, fmt)
        d.add_picture(str(img), width=Inches(0.4))
    p = tmp_path / "formats.docx"
    d.save(p)

    figs = extract_document(p).figures
    assert len(figs) == 3, "a figure was dropped, which would renumber the rest"
    for f in figs:
        assert f.mime in SUPPORTED_IMAGE_MIMES, f.mime
        assert f.deliverable
    assert figs[1].bytes.startswith(b"\x89PNG"), "BMP was not re-encoded"
    assert "converted from image/bmp" in figs[1].note


def test_an_unconvertible_figure_keeps_its_number(tmp_path):
    """Dropping it would make [Figure 3] point at what was really figure 4 —
    the exact mis-attribution the markers exist to prevent."""
    from scilink.parsers.docx_document import Figure, _normalize_figure

    bad = Figure(index=2, ext="x-emf", bytes=b"\x01\x00\x00\x00not-an-image")
    out = _normalize_figure(bad)
    assert out.index == 2, "renumbered"
    assert out.deliverable is False
    assert "could not be decoded" in out.note


def test_a_supported_mime_is_still_verified_not_trusted(tmp_path):
    """The declared type is a claim the bytes need not honour: a corrupt
    file named .png declares image/png and would pass a mime-only check on
    its way to being rejected by the provider."""
    from scilink.parsers.docx_document import Figure, _normalize_figure

    liar = Figure(index=0, ext="png",
                  bytes=b"\x89PNG\r\n\x1a\n" + b"\x00" * 200)
    out = _normalize_figure(liar)
    assert out.deliverable is False, "a corrupt PNG was trusted on its label"


def test_a_valid_supported_image_is_passed_through_untouched(tmp_path):
    """Verification must not re-encode what is already fine."""
    import io
    from PIL import Image
    from scilink.parsers.docx_document import Figure, _normalize_figure

    buf = io.BytesIO()
    Image.new("RGB", (20, 15), "teal").save(buf, "PNG")
    raw = buf.getvalue()
    out = _normalize_figure(Figure(index=0, ext="png", bytes=raw))
    assert out.bytes is raw or out.bytes == raw, "needlessly re-encoded"
    assert out.deliverable and out.note == ""


def test_a_figure_inside_a_table_cell_is_found(tmp_path):
    """Cells go through the same walker; an image in one must keep its place
    in the global numbering rather than being skipped."""
    from docx.shared import Inches
    from PIL import Image

    img = tmp_path / "c.png"
    Image.new("RGB", (30, 20), "navy").save(img)

    d = docx.Document()
    d.add_paragraph("Before:")
    d.add_picture(str(img), width=Inches(0.3))          # figure 1
    t = d.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "caption"
    t.rows[0].cells[1].paragraphs[0].add_run().add_picture(
        str(img), width=Inches(0.3))                    # figure 2, in a cell
    d.add_paragraph("After:")
    d.add_picture(str(img), width=Inches(0.3))          # figure 3
    p = tmp_path / "incell.docx"
    d.save(p)

    parsed = extract_document(p)
    composed = parsed.composed()
    assert len(parsed.figures) == 3, "the in-cell image was missed"
    # The in-cell marker belongs in the CELL, not the body text — that says
    # which cell holds the figure, which body placement could not.
    assert "| caption | [Figure 2] |" in composed, composed
    assert "[Figure 3]" in parsed.text, "numbering after the table broke"
    # Numbering must stay global: the in-cell image consumes slot 2, so the
    # paragraph image after the table is 3, not 2.
    assert parsed.text.index("[Figure 1]") < parsed.text.index("[Figure 3]")
