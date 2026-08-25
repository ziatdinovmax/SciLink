"""write_technical_document(style="memo"): one-page memo routing.

Offline: the author is stubbed. Checks the style flag reaches the author,
the length guard re-asks once and then ships with a warning, memo
revisions bypass the shrink guard, no workflow diagram is appended, and
the .docx twin renders the house layout.
"""
import json
from pathlib import Path

from docx import Document

from scilink.agents.planning_agents import orchestrator_tools as ot
from scilink.agents.planning_agents import planning_rag as ot_rag
from scilink.utils.md_to_docx import markdown_to_docx
from tests.test_edit_file_tool import make_tools

HEADER = ("*One-line subtitle*\n\n**Purpose:** Define X.\n**Scope:** Y only.\n"
          "**Date:** August 17, 2026\n\n**CORE RULE**  One sentence.")


def _memo_sections(n_words_body=60):
    body = " ".join(["word"] * n_words_body)
    return [{"heading": "", "body": HEADER},
            {"heading": "Scientific role", "body": f"**Role.** {body}"},
            {"heading": "Planning implication", "body": f"**Stage it.** {body}"}]


def _stub_author(monkeypatch, drafts):
    """Successive calls return successive drafts; records the kwargs."""
    calls = []

    def fake(request, kb_docs, model, generation_config, **kw):
        calls.append(kw)
        idx = min(len(calls) - 1, len(drafts) - 1)
        return {"sections": drafts[idx]}
    monkeypatch.setattr(ot, "author_technical_document", fake)
    return calls


def test_memo_style_reaches_author_and_exports_docx(tmp_path, monkeypatch):
    tools, cap, deleg = make_tools(tmp_path)
    tools._maybe_embed_workflow_diagram = lambda text, d, stem=None: text + "\nDIAGRAM\n"
    calls = _stub_author(monkeypatch, [_memo_sections()])

    out = json.loads(cap["write_technical_document"](
        request="one page on X", filename="x_memo.md", title="X memo",
        use_literature=False, style="memo"))

    assert out["status"] == "success"
    assert calls[0]["style"] == "memo"
    assert "TODAY" in (calls[0]["additional_context"] or "")
    assert out["style"] == "memo"
    md = Path(out["path"])
    assert "DIAGRAM" not in md.read_text(), "memos get no auto workflow diagram"
    docx = Path(out["docx"])
    assert docx.exists() and docx.suffix == ".docx"
    d = Document(docx)
    texts = [p.text for p in d.paragraphs]
    assert texts[0] == "TECHNICAL MEMO"
    assert texts[1] == "X memo"
    assert any(p.style.name == "Heading 1" and p.text == "Scientific role"
               for p in d.paragraphs)
    lead = next(p for p in d.paragraphs if p.text.startswith("Role."))
    assert lead.runs[0].bold and lead.runs[0].text == "Role."


def test_report_style_is_unchanged(tmp_path, monkeypatch):
    tools, cap, deleg = make_tools(tmp_path)
    tools._maybe_embed_workflow_diagram = lambda text, d, stem=None: text + "\nDIAGRAM\n"
    calls = _stub_author(monkeypatch, [_memo_sections()])
    out = json.loads(cap["write_technical_document"](
        request="a roadmap", filename="r.md", use_literature=False))
    assert out["status"] == "success" and out["style"] == "report"
    assert calls[0]["style"] == "report" and calls[0]["additional_context"] is None
    assert out["docx"] is None
    assert "DIAGRAM" in Path(out["path"]).read_text()


def _stub_condense(monkeypatch, result):
    calls = []

    def fake(draft, model, generation_config, **kw):
        calls.append((draft, kw))
        return result
    monkeypatch.setattr(ot_rag, "condense_memo", fake)
    return calls


def test_memo_length_guard_condenses_over_budget_draft(tmp_path, monkeypatch):
    tools, cap, deleg = make_tools(tmp_path)
    tools._maybe_embed_workflow_diagram = lambda text, d, stem=None: text
    calls = _stub_author(monkeypatch, [_memo_sections(500)])
    ccalls = _stub_condense(monkeypatch, {"sections": _memo_sections(50)})
    out = json.loads(cap["write_technical_document"](
        request="one page", filename="m.md", use_literature=False,
        style="memo"))
    assert len(calls) == 1, "the RAG author is not re-asked"
    assert len(ccalls) == 1 and "word word" in ccalls[0][0]
    assert ccalls[0][1]["max_words"] == 400
    assert out["status"] == "success" and "length_warning" not in out
    assert out["words"] < 200


def test_memo_condense_failure_ships_draft_with_warning(tmp_path, monkeypatch):
    tools, cap, deleg = make_tools(tmp_path)
    tools._maybe_embed_workflow_diagram = lambda text, d, stem=None: text
    _stub_author(monkeypatch, [_memo_sections(500)])
    _stub_condense(monkeypatch, {"error": "JSON Parsing Error"})
    out = json.loads(cap["write_technical_document"](
        request="one page", filename="m.md", use_literature=False,
        style="memo"))
    assert out["status"] == "success" and "length_warning" in out
    assert Path(out["path"]).exists() and Path(out["docx"]).exists()


def test_memo_condense_that_grew_is_discarded(tmp_path, monkeypatch):
    tools, cap, deleg = make_tools(tmp_path)
    tools._maybe_embed_workflow_diagram = lambda text, d, stem=None: text
    _stub_author(monkeypatch, [_memo_sections(500)])
    _stub_condense(monkeypatch, {"sections": _memo_sections(900)})
    out = json.loads(cap["write_technical_document"](
        request="one page", filename="m.md", use_literature=False,
        style="memo"))
    assert out["status"] == "success" and "length_warning" in out
    assert out["words"] < 1200  # the 500-word draft, not the 900-word one


def test_memo_revision_may_condense_a_long_document(tmp_path, monkeypatch):
    tools, cap, deleg = make_tools(tmp_path)
    tools._maybe_embed_workflow_diagram = lambda text, d, stem=None: text
    long_doc = tmp_path / "delegations" / "01_x" / "white_paper.md"
    long_doc.parent.mkdir(parents=True)
    long_doc.write_text("# White paper\n\n" + ("## S\n\n" + "prose " * 400 + "\n\n") * 5)
    calls = _stub_author(monkeypatch, [_memo_sections()])

    out = json.loads(cap["write_technical_document"](
        request="condense to a memo", revise_path=str(long_doc),
        use_literature=False, style="memo"))

    assert out["status"] == "success", out
    assert calls[0]["revise_document"] and calls[0]["style"] == "memo"
    assert out["words"] < 200 and out["revised_in_place"]
    assert Path(out["docx"]).exists()
    assert list(deleg.glob("white_paper.before_revision*"))


def test_report_revision_still_refuses_shrinkage(tmp_path, monkeypatch):
    tools, cap, deleg = make_tools(tmp_path)
    tools._maybe_embed_workflow_diagram = lambda text, d, stem=None: text
    long_doc = tmp_path / "delegations" / "01_x" / "white_paper.md"
    long_doc.parent.mkdir(parents=True)
    long_doc.write_text("# White paper\n\n" + ("## S\n\n" + "prose " * 400 + "\n\n") * 5)
    _stub_author(monkeypatch, [_memo_sections()])
    out = json.loads(cap["write_technical_document"](
        request="tighten", revise_path=str(long_doc), use_literature=False))
    assert out["status"] == "error" and "Revision aborted" in out["message"]


def test_docx_converter_tables_bullets_and_header_lines(tmp_path):
    md = tmp_path / "m.md"
    md.write_text("# T\n\n**Purpose:** p.\n\n**CORE UNIT**  u.\n\n## A\n\n"
                  "- one\n- two\n\n| Track | Out |\n|---|---|\n| 1 | traj |\n")
    out = markdown_to_docx(md)
    d = Document(out)
    styles = [p.style.name for p in d.paragraphs]
    assert styles.count("List Bullet") == 2
    assert len(d.tables) == 1 and d.tables[0].cell(1, 1).text == "traj"
    purpose = next(p for p in d.paragraphs if p.text.startswith("Purpose:"))
    assert purpose.runs[0].bold


def test_memo_condense_runs_a_second_pass_when_still_long(tmp_path, monkeypatch):
    tools, cap, deleg = make_tools(tmp_path)
    tools._maybe_embed_workflow_diagram = lambda text, d, stem=None: text
    _stub_author(monkeypatch, [_memo_sections(500)])
    outs = [{"sections": _memo_sections(350)}, {"sections": _memo_sections(120)}]
    calls = []

    def fake(draft, model, generation_config, **kw):
        calls.append(draft)
        return outs[min(len(calls) - 1, 1)]
    monkeypatch.setattr(ot_rag, "condense_memo", fake)
    out = json.loads(cap["write_technical_document"](
        request="one page", filename="m.md", use_literature=False,
        style="memo"))
    assert len(calls) == 2
    assert out["words"] < 400 and "length_warning" not in out
